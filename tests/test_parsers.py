"""
Unit tests for DocumentParser.
"""
import io
import pytest
from src.parsers.document_parser import DocumentParser
import docx


def test_parse_plain_text():
    sample_text = "This is a plain text resume content for testing."
    result = DocumentParser.extract_text(sample_text.encode("utf-8"), filename="test.txt")
    assert "plain text resume" in result


def test_parse_docx():
    doc = docx.Document()
    doc.add_heading("Resume of John Developer", 0)
    doc.add_paragraph("Skills: Python, Django, SQL")
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    
    result = DocumentParser.extract_text(buf.getvalue(), filename="test.docx")
    assert "John Developer" in result
    assert "Python" in result


def test_invalid_extension():
    content = b"Some raw bytes"
    result = DocumentParser.extract_text(content, filename="test.xyz")
    assert "Some raw bytes" in result
