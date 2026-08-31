"""
Multi-format Document Parser for PDF, DOCX, and TXT Resumes.
Supports direct file paths and raw byte streams (for Web UI / API uploads).
"""
import io
import os
import re
from pathlib import Path
from typing import Union, BinaryIO, Optional
import pypdf
import docx


class DocumentParser:
    """
    Parses documents of various formats (PDF, DOCX, TXT) into clean raw text.
    """

    @staticmethod
    def extract_text(
        source: Union[str, Path, bytes, BinaryIO], 
        filename: Optional[str] = None
    ) -> str:
        """
        Extract text from file path or bytes buffer.
        
        Args:
            source: File path (str/Path) or bytes / BinaryIO stream.
            filename: Optional filename hint (used to detect format if source is bytes).
            
        Returns:
            Extracted clean plain text.
        """
        if isinstance(source, (str, Path)):
            path = Path(source)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {path}")
            
            ext = path.suffix.lower()
            with open(path, "rb") as f:
                content_bytes = f.read()
            return DocumentParser._parse_by_extension(content_bytes, ext, filename=path.name)
            
        elif isinstance(source, bytes):
            ext = Path(filename).suffix.lower() if filename else ".txt"
            return DocumentParser._parse_by_extension(source, ext, filename=filename or "document")
            
        elif hasattr(source, "read"):
            content_bytes = source.read()
            ext = Path(filename).suffix.lower() if filename else ".txt"
            return DocumentParser._parse_by_extension(content_bytes, ext, filename=filename or "document")
            
        else:
            raise ValueError(f"Unsupported source type: {type(source)}")

    @staticmethod
    def _parse_by_extension(content_bytes: bytes, ext: str, filename: str) -> str:
        """Dispatches extraction based on file extension."""
        if ext == ".pdf":
            return DocumentParser._extract_from_pdf(content_bytes)
        elif ext in [".docx", ".doc"]:
            return DocumentParser._extract_from_docx(content_bytes)
        elif ext in [".txt", ".md", ".rtf", ""]:
            return DocumentParser._extract_from_txt(content_bytes)
        else:
            # Fallback: attempt utf-8 text decode
            try:
                return content_bytes.decode("utf-8", errors="ignore")
            except Exception as e:
                raise ValueError(f"Unsupported file format '{ext}' for file '{filename}': {e}")

    @staticmethod
    def _extract_from_pdf(content_bytes: bytes) -> str:
        """Extracts text from PDF bytes using PyPDF."""
        text_parts = []
        try:
            reader = pypdf.PdfReader(io.BytesIO(content_bytes))
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            extracted = "\n".join(text_parts)
            return DocumentParser._normalize_newlines(extracted)
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from PDF: {str(e)}")

    @staticmethod
    def _extract_from_docx(content_bytes: bytes) -> str:
        """Extracts text from DOCX bytes using python-docx."""
        try:
            doc = docx.Document(io.BytesIO(content_bytes))
            text_parts = []
            
            # Paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
                    
            # Tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
                        
            extracted = "\n".join(text_parts)
            return DocumentParser._normalize_newlines(extracted)
        except Exception as e:
            raise RuntimeError(f"Failed to extract text from DOCX: {str(e)}")

    @staticmethod
    def _extract_from_txt(content_bytes: bytes) -> str:
        """Extracts text from plain text bytes with multiple encoding attempts."""
        encodings = ["utf-8", "latin-1", "utf-16", "cp1252"]
        for enc in encodings:
            try:
                decoded = content_bytes.decode(enc)
                return DocumentParser._normalize_newlines(decoded)
            except UnicodeDecodeError:
                continue
        # Fallback with ignore
        decoded = content_bytes.decode("utf-8", errors="ignore")
        return DocumentParser._normalize_newlines(decoded)

    @staticmethod
    def _normalize_newlines(text: str) -> str:
        """Standardizes line breaks and whitespace formatting."""
        if not text:
            return ""
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Remove consecutive blank lines
        text = re.sub(r"\n\s*\n", "\n\n", text)
        return text.strip()
